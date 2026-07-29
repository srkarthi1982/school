from __future__ import annotations

import io
import logging
import re
from typing import Type

from fastapi import HTTPException, status
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.modules.course_info.models import (
    CourseInfoEntryStandard,
    CourseInfoGeneralInformation,
    CourseInfoInstructorQualificationRequirement,
    CourseInfoLessonCreation,
    CourseInfoLessonCreationLesson,
    CourseInfoLessonCreationLessonConduct,
    CourseInfoLessonCreationLessonResource,
    CourseInfoLessonCreationLessonUnit,
    CourseInfoLessonPlanning,
    CourseInfoLessonPlanningClassificationOfPeriod,
    CourseInfoLessonPlanningEnablingObjective,
    CourseInfoLessonPlanningEnvironment,
    CourseInfoLessonPlanningTeachingPoint,
    CourseInfoLessonPlanningTrainingObjective,
    CourseInfoLessonPlanningTypeOfPeriod,
    CourseInfoPersonnelRequirement,
    CourseInfoResources,
    CourseInfoResourcesAircraftItem,
    CourseInfoResourcesAircraftArmamentItem,
    CourseInfoResourcesAircraftMissionEquipmentItem,
    CourseInfoResourcesAviationLifeSupportItem,
    CourseInfoResourcesClassroomRequirementItem,
    CourseInfoResourcesGroundArmamentItem,
    CourseInfoResourcesGroundMaintenanceItem,
    CourseInfoResourcesMissionPlanningSystemItem,
    CourseInfoResourcesPersonalFlightEquipmentItem,
    CourseInfoResourcesSimulatorTypeItem,
    CourseInfoResourcesTrainingMaterialAidItem,
    MasterAircraftArmament,
    MasterAircraftMissionEquipment,
    MasterAircraftType,
    MasterAviationLifeSupportEquipment,
    MasterClassroomRequirement,
    MasterCourseEntryStandard,
    MasterEnablingObjective,
    MasterEnvironment,
    MasterGroundArmament,
    MasterGroundMaintenanceEquipment,
    MasterMissionPlanningSystem,
    MasterPeriodClassification,
    MasterPeriodType,
    MasterPersonalFlightEquipment,
    MasterSimulatorType,
    MasterInstructorQualificationRequirement,
    MasterTeachingPoint,
    MasterTrainingMaterialAid,
    MasterTrainingObjective,
)
from app.modules.course_info.schemas import (
    AircraftItem,
    AircraftArmamentItem,
    AircraftMissionEquipmentItem,
    AviationLifeSupportItem,
    ClassroomRequirementItem,
    ClassificationOfPeriodItem,
    CourseEntryStandardItem,
    CourseInfoDetailResponse,
    CourseInfoTabSummary,
    EnablingObjectiveItem,
    EnvironmentItem,
    GeneralInformationData,
    GroundArmamentItem,
    GroundMaintenanceItem,
    InstructorQualificationRequirementItem,
    LessonAssociationsResponse,
    LessonConductItem,
    LessonCreationData,
    LessonItem,
    LessonResourceCategoryOption,
    LessonResourceItem,
    LessonUnitItem,
    LessonPlanningData,
    ResourceOptionResponse,
    MissionPlanningSystemItem,
    PersonnelRequirementData,
    PersonalFlightEquipmentItem,
    ProgrammedWorkingHours,
    ResourcesData,
    SimulatorTypeItem,
    TeachingPointItem,
    TpLinkEnablingObjective,
    TpLinkPreviewData,
    TpLinkTeachingPoint,
    TpLinkTrainingObjective,
    TrainingMaterialAidItem,
    TrainingObjectiveItem,
    TypeOfPeriodItem,
)
from app.modules.course_master.models import CourseMaster

logger = logging.getLogger(__name__)


# Combo-box ``kind`` -> per-field master model. The keys are the only kinds
# the option endpoints accept; anything else is rejected as a 400.
KIND_TO_MASTER_MODEL: dict[str, Type] = {
    # Resources tab
    "aircraft_type": MasterAircraftType,
    "aircraft_mission_equipment": MasterAircraftMissionEquipment,
    "aircraft_armament": MasterAircraftArmament,
    "ground_equipment_simulator_type": MasterSimulatorType,
    "ground_equipment_mission_planning_system": MasterMissionPlanningSystem,
    "ground_equipment_maintenance": MasterGroundMaintenanceEquipment,
    "ground_equipment_armament": MasterGroundArmament,
    "personal_flight_equipment": MasterPersonalFlightEquipment,
    "aviation_life_support_equipment": MasterAviationLifeSupportEquipment,
    "classroom_requirements": MasterClassroomRequirement,
    "training_material_aids": MasterTrainingMaterialAid,
    # General Information tab
    "course_entry_standard": MasterCourseEntryStandard,
    # Lesson Planning tab
    "type_of_period": MasterPeriodType,
    "classification_of_period": MasterPeriodClassification,
    "environment": MasterEnvironment,
    "training_objective": MasterTrainingObjective,
    "enabling_objective": MasterEnablingObjective,
    "teaching_point": MasterTeachingPoint,
    # Personnel Requirement tab
    "instructor_qualification_requirement": MasterInstructorQualificationRequirement,
}


# Mirrors the ``implemented: true`` flags in the frontend catalogue
# (frontend/.../course-info/editor/catalogue.ts). Tabs without a working editor
# don't contribute to either side of the percentage — completing every
# implemented tab should read 100%.
IMPLEMENTED_TAB_NOS: frozenset[str] = frozenset(
    {"general", "personnel_requirement", "resources", "lesson_planning", "lesson_creation"}
)

TOTAL_TAB_COUNT = len(IMPLEMENTED_TAB_NOS)

# Catalogue used to seed display names. Kept in sync with the frontend
# `TABS` array.
TAB_CATALOGUE: tuple[tuple[str, str], ...] = (
    ("general", "General Information"),
    ("resources", "Resources"),
    ("lesson_planning", "Lesson Planning"),
    ("lesson_creation", "Lesson Creation"),
    ("personnel_requirement", "Personnel Requirement"),
)

# Maps a tab slug -> the relationship attribute on ``CourseMaster`` holding the
# typed row for that tab. Each future tab gets its own typed table and an
# entry here. Untyped/unimplemented tabs are absent.
TAB_RELATIONSHIP_NAMES: dict[str, str] = {
    "general": "general",
    "personnel_requirement": "personnel_requirement",
    "resources": "resources",
    "lesson_planning": "lesson_planning",
    "lesson_creation": "lesson_creation",
}


def _get_tab_row(master: CourseMaster, tab_no: str):
    attr = TAB_RELATIONSHIP_NAMES.get(tab_no)
    if attr is None:
        return None
    return getattr(master, attr, None)


def compute_completion_pct(master: CourseMaster) -> float:
    """Percentage of implemented tabs that are marked complete."""
    if TOTAL_TAB_COUNT == 0:
        return 0.0
    completed = 0
    for tab_no in IMPLEMENTED_TAB_NOS:
        row = _get_tab_row(master, tab_no)
        if row is not None and getattr(row, "status", None) == "complete":
            completed += 1
    return round((completed / TOTAL_TAB_COUNT) * 100, 2)


def get_or_raise_course_master(db: Session, course_master_id: int) -> CourseMaster:
    master = db.get(CourseMaster, course_master_id)
    if not master:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Course information not found"
        )
    return master


def build_tab_summaries(master: CourseMaster) -> list[CourseInfoTabSummary]:
    """Return one summary per tab in the catalogue, even if no row exists yet."""
    summaries: list[CourseInfoTabSummary] = []
    for slug, name in TAB_CATALOGUE:
        row = _get_tab_row(master, slug)
        summaries.append(
            CourseInfoTabSummary(
                tab_no=slug,
                tab_name=name,
                status=getattr(row, "status", "incomplete") if row else "incomplete",
            )
        )
    return summaries


def serialize_course_info_detail(master: CourseMaster) -> CourseInfoDetailResponse:
    return CourseInfoDetailResponse(
        id=master.id,
        course_master_id=master.id,
        course_title=master.title,
        version=master.version,
        created_by_id=master.created_by_id,
        updated_by_id=master.updated_by_id,
        created_at=master.created_at,
        updated_at=master.updated_at,
        completion_pct=compute_completion_pct(master),
        course_master_status=master.status,
        tabs=build_tab_summaries(master),
        currencies_certificate_completion=master.currencies_certificate_completion,
        flight_package_completion=getattr(master, "flight_package_completion", 0),
        task_association_completion=getattr(master, "task_association_completion", 0),
        flight_pack_association_completion=getattr(master, "flight_pack_association_completion", 0),
    )


def _guard_unmark_when_approved(master: CourseMaster, row, new_status: str) -> None:
    """Once the course master is approved, completed tabs are locked: they
    can no longer be moved back to ``draft``/``incomplete``."""
    if (
        row is not None
        and row.status == "complete"
        and new_status != "complete"
        and master.status == "approved"
    ):
        raise HTTPException(
            status_code=409,
            detail="Course is approved; tabs cannot be unmarked complete.",
        )


# ─── General Information tab ────────────────────────────────────────────────

def _general_to_data(row: CourseInfoGeneralInformation, master: CourseMaster) -> GeneralInformationData:
    return GeneralInformationData(
        courseTitle=master.title,
        courseDuration=row.course_duration,
        programmedWorkingHours=ProgrammedWorkingHours(
            frequency=row.programmed_working_frequency,
            startTime=row.programmed_working_start_time,
            endTime=row.programmed_working_end_time,
        ),
        ctpVersion=master.ctp_version,
        courseAim=row.course_aim,
        courseEntryStandard=[
            CourseEntryStandardItem(courseEntryStandardId=es.course_entry_standard_id)
            for es in row.entry_standards
        ],
    )


def _rebuild_label_only_item_list(
    db: Session,
    *,
    collection,
    items,
    master_model: Type,
    item_orm_cls: Type,
    fk_attr: str,
    id_attr: str,
) -> None:
    """Replace an ordered child collection with a cleaned, deduped list.

    Skips rows with a blank FK id, dedupes by FK id, validates every FK id
    exists in the master table, and re-appends rows with sequential
    ``order_index``.  Variant of ``_rebuild_item_list`` for fields that have
    no quantity column.
    """
    cleaned: list[int] = []
    seen: set[int] = set()
    for it in items:
        fk = getattr(it, id_attr)
        if fk is None:
            continue
        if fk in seen:
            continue
        seen.add(fk)
        _ensure_master_row_exists(db, master_model, fk)
        cleaned.append(fk)
    collection.clear()
    for idx, fk in enumerate(cleaned):
        collection.append(item_orm_cls(order_index=idx, **{fk_attr: fk}))


def _apply_general_data(
    db: Session, row: CourseInfoGeneralInformation, data: GeneralInformationData,
    master: CourseMaster,
) -> None:
    if data.courseTitle:
        master.title = data.courseTitle
    row.course_duration = data.courseDuration
    pwh = data.programmedWorkingHours
    if pwh is not None:
        row.programmed_working_frequency = pwh.frequency
        row.programmed_working_start_time = pwh.startTime
        row.programmed_working_end_time = pwh.endTime
    else:
        row.programmed_working_frequency = None
        row.programmed_working_start_time = None
        row.programmed_working_end_time = None
    master.ctp_version = data.ctpVersion
    row.course_aim = data.courseAim
    _rebuild_label_only_item_list(
        db,
        collection=row.entry_standards,
        items=data.courseEntryStandard or [],
        master_model=MasterCourseEntryStandard,
        item_orm_cls=CourseInfoEntryStandard,
        fk_attr="course_entry_standard_id",
        id_attr="courseEntryStandardId",
    )


def get_general_tab(master: CourseMaster) -> tuple[str, GeneralInformationData | None]:
    """Return ``(status, data)`` for the General Information tab."""
    row = master.general
    if row is None:
        return "incomplete", GeneralInformationData(courseTitle=master.title)
    return row.status, _general_to_data(row, master)


def upsert_general_tab(
    db: Session,
    master: CourseMaster,
    new_status: str,
    data: GeneralInformationData,
    user_id: int,
) -> CourseInfoGeneralInformation:
    if new_status not in ("incomplete", "draft", "complete"):
        raise HTTPException(
            status_code=400,
            detail="status must be 'incomplete', 'draft' or 'complete'",
        )

    row = master.general
    _guard_unmark_when_approved(master, row, new_status)
    if row is None:
        row = CourseInfoGeneralInformation(
            status=new_status,
            created_by_id=user_id,
            updated_by_id=user_id,
        )
        _apply_general_data(db, row, data, master)
        # Use the relationship so master.general reflects the new row
        # before we compute the completion percentage below.
        master.general = row
    else:
        row.status = new_status
        _apply_general_data(db, row, data, master)
        row.updated_by_id = user_id

    _sync_master_completion(master)

    db.commit()
    db.refresh(row)
    return row


# ─── Resources tab ──────────────────────────────────────────────────────────

def _resources_to_data(row: CourseInfoResources) -> ResourcesData:
    return ResourcesData(
        aircraftItems=[
            AircraftItem(
                aircraftTypeId=item.aircraft_type_id,
                numberOfAircraft=item.number_of_aircraft,
            )
            for item in row.aircraft_items
        ],
        aircraftMissionEquipmentItems=[
            AircraftMissionEquipmentItem(
                aircraftMissionEquipmentId=item.aircraft_mission_equipment_id,
                quantity=item.quantity,
            )
            for item in row.aircraft_mission_equipment_items
        ],
        aircraftArmamentItems=[
            AircraftArmamentItem(
                aircraftArmamentId=item.aircraft_armament_id,
                quantity=item.quantity,
            )
            for item in row.aircraft_armament_items
        ],
        groundEquipmentSimulatorTypeItems=[
            SimulatorTypeItem(
                simulatorTypeId=item.simulator_type_id,
                quantity=item.quantity,
            )
            for item in row.ground_equipment_simulator_type_items
        ],
        groundEquipmentMissionPlanningSystemItems=[
            MissionPlanningSystemItem(
                missionPlanningSystemId=item.mission_planning_system_id,
                quantity=item.quantity,
            )
            for item in row.ground_equipment_mission_planning_system_items
        ],
        groundEquipmentMaintenanceItems=[
            GroundMaintenanceItem(
                groundMaintenanceId=item.ground_maintenance_equipment_id,
                quantity=item.quantity,
            )
            for item in row.ground_equipment_maintenance_items
        ],
        groundEquipmentArmamentItems=[
            GroundArmamentItem(
                groundArmamentId=item.ground_armament_id,
                quantity=item.quantity,
            )
            for item in row.ground_equipment_armament_items
        ],
        personalFlightEquipmentItems=[
            PersonalFlightEquipmentItem(
                personalFlightEquipmentId=item.personal_flight_equipment_id,
                quantity=item.quantity,
            )
            for item in row.personal_flight_equipment_items
        ],
        aviationLifeSupportEquipmentItems=[
            AviationLifeSupportItem(
                aviationLifeSupportEquipmentId=item.aviation_life_support_equipment_id,
                quantity=item.quantity,
            )
            for item in row.aviation_life_support_equipment_items
        ],
        classroomRequirementItems=[
            ClassroomRequirementItem(
                classroomRequirementId=item.classroom_requirement_id,
                quantity=item.quantity,
            )
            for item in row.classroom_requirement_items
        ],
        trainingMaterialAidItems=[
            TrainingMaterialAidItem(
                trainingMaterialAidId=item.training_material_aid_id,
                quantity=item.quantity,
            )
            for item in row.training_material_aid_items
        ],
    )


def _rebuild_item_list(
    db: Session,
    *,
    collection,
    items,
    master_model: Type,
    item_orm_cls: Type,
    fk_attr: str,
    id_attr: str,
    quantity_attr: str,
    orm_quantity_attr: str,
) -> None:
    """Replace an ordered child collection with a cleaned, deduped list.

    Skips fully-blank rows, dedupes by FK id, validates every FK id exists in
    the master table, and re-appends rows with sequential ``order_index``.
    """
    cleaned = []
    seen: set[int] = set()
    for it in items:
        fk = getattr(it, id_attr)
        qty = getattr(it, quantity_attr)
        if fk is None and qty is None:
            continue
        if fk is not None:
            if fk in seen:
                continue
            seen.add(fk)
        _ensure_master_row_exists(db, master_model, fk)
        cleaned.append(it)
    collection.clear()
    for idx, it in enumerate(cleaned):
        collection.append(
            item_orm_cls(
                order_index=idx,
                **{fk_attr: getattr(it, id_attr), orm_quantity_attr: getattr(it, quantity_attr)},
            )
        )


# (schema_field_name, orm_relationship_attr, master_model, item_orm_cls,
#  item_fk_attr, item_id_attr, item_quantity_attr, item_orm_quantity_attr)
_RESOURCES_ITEM_FIELDS: tuple[tuple[str, str, Type, Type, str, str, str, str], ...] = (
    ("aircraftItems", "aircraft_items", MasterAircraftType, CourseInfoResourcesAircraftItem, "aircraft_type_id", "aircraftTypeId", "numberOfAircraft", "number_of_aircraft"),
    ("aircraftMissionEquipmentItems", "aircraft_mission_equipment_items", MasterAircraftMissionEquipment, CourseInfoResourcesAircraftMissionEquipmentItem, "aircraft_mission_equipment_id", "aircraftMissionEquipmentId", "quantity", "quantity"),
    ("aircraftArmamentItems", "aircraft_armament_items", MasterAircraftArmament, CourseInfoResourcesAircraftArmamentItem, "aircraft_armament_id", "aircraftArmamentId", "quantity", "quantity"),
    ("groundEquipmentSimulatorTypeItems", "ground_equipment_simulator_type_items", MasterSimulatorType, CourseInfoResourcesSimulatorTypeItem, "simulator_type_id", "simulatorTypeId", "quantity", "quantity"),
    ("groundEquipmentMissionPlanningSystemItems", "ground_equipment_mission_planning_system_items", MasterMissionPlanningSystem, CourseInfoResourcesMissionPlanningSystemItem, "mission_planning_system_id", "missionPlanningSystemId", "quantity", "quantity"),
    ("groundEquipmentMaintenanceItems", "ground_equipment_maintenance_items", MasterGroundMaintenanceEquipment, CourseInfoResourcesGroundMaintenanceItem, "ground_maintenance_equipment_id", "groundMaintenanceId", "quantity", "quantity"),
    ("groundEquipmentArmamentItems", "ground_equipment_armament_items", MasterGroundArmament, CourseInfoResourcesGroundArmamentItem, "ground_armament_id", "groundArmamentId", "quantity", "quantity"),
    ("personalFlightEquipmentItems", "personal_flight_equipment_items", MasterPersonalFlightEquipment, CourseInfoResourcesPersonalFlightEquipmentItem, "personal_flight_equipment_id", "personalFlightEquipmentId", "quantity", "quantity"),
    ("aviationLifeSupportEquipmentItems", "aviation_life_support_equipment_items", MasterAviationLifeSupportEquipment, CourseInfoResourcesAviationLifeSupportItem, "aviation_life_support_equipment_id", "aviationLifeSupportEquipmentId", "quantity", "quantity"),
    ("classroomRequirementItems", "classroom_requirement_items", MasterClassroomRequirement, CourseInfoResourcesClassroomRequirementItem, "classroom_requirement_id", "classroomRequirementId", "quantity", "quantity"),
    ("trainingMaterialAidItems", "training_material_aid_items", MasterTrainingMaterialAid, CourseInfoResourcesTrainingMaterialAidItem, "training_material_aid_id", "trainingMaterialAidId", "quantity", "quantity"),
)


def _apply_resources_data(
    db: Session, row, data: ResourcesData, *, fields=_RESOURCES_ITEM_FIELDS
) -> None:
    # ``fields`` defaults to the master child ORM classes; the Course Selection
    # instance category passes its own mirror classes (same shape) to reuse this.
    for schema_field, orm_rel, master_model, item_orm_cls, fk_attr, id_attr, qty_attr, orm_qty_attr in fields:
        item_list = getattr(data, schema_field)
        _rebuild_item_list(
            db,
            collection=getattr(row, orm_rel),
            items=item_list,
            master_model=master_model,
            item_orm_cls=item_orm_cls,
            fk_attr=fk_attr,
            id_attr=id_attr,
            quantity_attr=qty_attr,
            orm_quantity_attr=orm_qty_attr,
        )


def get_resources_tab(master: CourseMaster) -> tuple[str, ResourcesData | None]:
    """Return ``(status, data)`` for the Resources tab."""
    row = master.resources
    if row is None:
        return "incomplete", None
    return row.status, _resources_to_data(row)


def upsert_resources_tab(
    db: Session,
    master: CourseMaster,
    new_status: str,
    data: ResourcesData,
    user_id: int,
) -> CourseInfoResources:
    if new_status not in ("incomplete", "draft", "complete"):
        raise HTTPException(
            status_code=400,
            detail="status must be 'incomplete', 'draft' or 'complete'",
        )

    row = master.resources
    _guard_unmark_when_approved(master, row, new_status)
    if row is None:
        row = CourseInfoResources(
            status=new_status,
            created_by_id=user_id,
            updated_by_id=user_id,
        )
        _apply_resources_data(db, row, data)
        master.resources = row
    else:
        row.status = new_status
        _apply_resources_data(db, row, data)
        row.updated_by_id = user_id

    _sync_master_completion(master)

    db.commit()
    db.refresh(row)
    return row


# ─── Combo-box master option dictionaries ───────────────────────────────────

def _resolve_master_model(kind: str) -> Type:
    model = KIND_TO_MASTER_MODEL.get(kind)
    if model is None:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown resource option kind: {kind!r}",
        )
    return model


def _ensure_master_row_exists(db: Session, model: Type, row_id: int | None) -> None:
    """Validate that ``row_id`` actually points at a row in ``model``.

    Caller may pass ``None`` for "no selection" — that always passes.
    """
    if row_id is None:
        return
    exists = db.scalar(select(model.id).where(model.id == row_id))
    if exists is None:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown {model.__tablename__} id: {row_id}",
        )


def list_resource_options(db: Session, kind: str) -> list:
    """Return every row of the master table backing ``kind``, sorted by label."""
    model = _resolve_master_model(kind)
    stmt = select(model).order_by(model.label)
    return list(db.scalars(stmt).all())


def create_resource_option(
    db: Session, kind: str, label: str, user_id: int
):
    """Insert a new label into the master table backing ``kind``.

    Idempotent on the ``(table, label)`` pair — if the label already exists,
    the existing row is returned rather than creating a duplicate.
    """
    model = _resolve_master_model(kind)
    clean = label.strip()
    if not clean:
        raise HTTPException(status_code=400, detail="label must not be empty")

    existing = db.scalar(select(model).where(model.label == clean))
    if existing is not None:
        return existing

    option = model(label=clean, created_by_id=user_id)
    db.add(option)
    db.commit()
    db.refresh(option)
    return option


def bulk_get_or_create_options(
    db: Session, kind: str, labels: list[str], user_id: int
) -> dict[str, int]:
    """Resolve many labels to option ids in one round-trip, creating any missing.

    Returns a ``{cleaned_label: id}`` map. Used by .docx imports that would
    otherwise call :func:`create_resource_option` hundreds of times (one commit
    each) — here it is a single SELECT plus one INSERT batch.
    """
    model = _resolve_master_model(kind)
    cleaned: list[str] = []
    seen: set[str] = set()
    for raw in labels:
        c = (raw or "").strip()
        if c and c not in seen:
            seen.add(c)
            cleaned.append(c)
    if not cleaned:
        return {}

    id_by_label: dict[str, int] = {
        o.label: o.id
        for o in db.scalars(select(model).where(model.label.in_(cleaned))).all()
    }
    new_rows = [
        model(label=c, created_by_id=user_id) for c in cleaned if c not in id_by_label
    ]
    if new_rows:
        db.add_all(new_rows)
        db.flush()  # populate PKs without expiring the instances
        for r in new_rows:
            id_by_label[r.label] = r.id
        db.commit()
    return id_by_label


# ─── General Information .docx import ────────────────────────────────────────
# Mirrors the frequency labels offered by the General Information tab on the
# frontend (frontend/.../tabs/GeneralInformationTab.tsx). Note the en-dash in
# "Mon–Fri".
_FREQUENCY_OPTIONS: tuple[str, ...] = ("Daily", "Weekly", "Mon–Fri", "Weekends", "Custom")

# Lines inside the entry-standard cell that are headers/intros rather than
# individual pre-requisites — dropped before auto-creating options.
_ENTRY_STANDARD_INTRO_PREFIXES: tuple[str, ...] = (
    "candidates attending this course must have",
    "candidates must have",
    "candidates must",
)


def _normalize_label(text: str) -> str:
    """Lowercase + collapse whitespace so doc labels match tolerantly.

    The source documents are inconsistent (e.g. ``"Program med Working Hours"``),
    so we strip every run of whitespace down to a single space.
    """
    return re.sub(r"\s+", " ", text or "").strip().lower()


def _parse_duration_days(text: str) -> int | None:
    """Extract the training-day count, e.g. ``"22 Weeks (109 training days)"`` → 109."""
    m = re.search(r"(\d+)\s*training\s*day", text, re.I)
    return int(m.group(1)) if m else None


def _parse_working_hours(text: str) -> ProgrammedWorkingHours:
    """Parse e.g. ``"Daily from 0700 to 1430"`` into frequency + start/end times."""
    frequency: str | None = None
    norm = _normalize_label(text)
    for opt in _FREQUENCY_OPTIONS:
        if opt.lower() in norm:
            frequency = opt
            break

    times: list[str] = []
    for hh, mm in re.findall(r"\b(\d{1,2})[:.]?(\d{2})\b", text):
        hour, minute = int(hh), int(mm)
        if 0 <= hour <= 23 and 0 <= minute <= 59:
            times.append(f"{hour:02d}:{minute:02d}")

    return ProgrammedWorkingHours(
        frequency=frequency,
        startTime=times[0] if len(times) >= 1 else None,
        endTime=times[1] if len(times) >= 2 else None,
    )


def parse_general_information_docx(
    db: Session, file_bytes: bytes, user_id: int
) -> GeneralInformationData:
    """Parse a General Course Information .docx into General Information form data.

    The source document is a two-column key/value table. We match the label
    column tolerantly and map only the fields the form actually has; free-text
    entry standards are turned into ``course_entry_standard`` master options via
    the idempotent :func:`create_resource_option`.
    """
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    try:
        import docx  # python-docx; imported lazily so the rest of the module
                     # loads even if the optional dependency is missing.

        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001 — any failure means a bad upload
        logger.warning(
            "General Information .docx import failed to parse (%d bytes): %s",
            len(file_bytes),
            exc,
        )
        raise HTTPException(
            status_code=400,
            detail="Could not read .docx file. Make sure it is a valid Word document (.docx, not .doc).",
        ) from exc

    # Build {normalized_label: (raw_value_text, [value_lines])} from every
    # two-column row across all tables. First column wins on duplicate labels.
    fields: dict[str, tuple[str, list[str]]] = {}
    for table in document.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = _normalize_label(cells[0].text)
            if not label or label in fields:
                continue
            value_cell = cells[1]
            lines = [p.text.strip() for p in value_cell.paragraphs if p.text.strip()]
            fields[label] = (value_cell.text.strip(), lines)

    def value_of(*label_substrings: str) -> tuple[str, list[str]] | None:
        for norm_label, payload in fields.items():
            if all(sub in norm_label for sub in label_substrings):
                return payload
        return None

    title = value_of("course title")
    aim = value_of("course aim")
    duration = value_of("course duration")
    working = value_of("working hours")
    entry = value_of("entry standard") or value_of("pre-requisite") or value_of("prerequisite")

    course_title = title[0] if title else None
    # Join multi-paragraph aim cells into a single block of text.
    course_aim = "\n".join(aim[1]) if aim and aim[1] else (aim[0] if aim else None)
    course_duration = _parse_duration_days(duration[0]) if duration else None
    working_hours = _parse_working_hours(working[0]) if working else ProgrammedWorkingHours()

    entry_items: list[CourseEntryStandardItem] = []
    if entry:
        for line in entry[1]:
            norm = _normalize_label(line)
            if any(norm.startswith(pfx) for pfx in _ENTRY_STANDARD_INTRO_PREFIXES):
                continue
            option = create_resource_option(db, "course_entry_standard", line, user_id)
            entry_items.append(
                CourseEntryStandardItem(courseEntryStandardId=option.id)
            )

    return GeneralInformationData(
        courseTitle=course_title,
        courseDuration=course_duration,
        programmedWorkingHours=working_hours,
        courseAim=course_aim,
        courseEntryStandard=entry_items,
    )


# ─── Resources .docx import ─────────────────────────────────────────────────
# Values that mean "nothing required" — skipped rather than turned into items.
_SKIP_RESOURCE_VALUES: frozenset[str] = frozenset(
    {"nil", "none", "n/a", "na", "not required", "not applicable", "-", "—"}
)


def _parse_resource_item_line(line: str) -> tuple[str, int | None] | None:
    """Turn one resource line into ``(label, quantity)`` or ``None`` to skip.

    Handles leading counts like ``"4 x MH-60M"`` / ``"8 × Copies"`` → quantity 4/8
    with the remainder as the label. Sub-headers ending in ``:`` and "nil"-style
    placeholders are dropped.
    """
    s = (line or "").strip().lstrip("•-*").strip()
    if not s or s.lower() in _SKIP_RESOURCE_VALUES or s.endswith(":"):
        return None
    m = re.match(r"^(\d+)\s*[x×]\s*(.+)$", s, re.I)
    if m:
        return m.group(2).strip(), int(m.group(1))
    return s, None


def _split_resource_cell(cell_lines: list[str]) -> list[str]:
    """Split each cell line into individual items.

    Comma-separated lists are only split when *every* part looks like a
    quantified item (e.g. ``"4 x MH-60M, 2 x Spare MH-60M"``) so prose
    containing commas (e.g. ``"Stationary (Pens, markers)"``) stays intact.
    """
    items: list[str] = []
    for line in cell_lines:
        parts = [p.strip() for p in line.split(",")]
        if len(parts) > 1 and all(
            re.match(r"^\s*\d+\s*[x×]\s*.+", p, re.I) for p in parts
        ):
            items.extend(parts)
        else:
            items.append(line)
    return items


def parse_resources_docx(
    db: Session, file_bytes: bytes, user_id: int
) -> ResourcesData:
    """Parse a Course Resources .docx into Resources form data.

    The source document is a two-column key/value table. Only the categories the
    Resources tab actually has are mapped; each item is turned into the matching
    master option via the idempotent :func:`create_resource_option`.
    """
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    try:
        import docx  # python-docx, imported lazily.

        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001 — any failure means a bad upload
        logger.warning(
            "Resources .docx import failed to parse (%d bytes): %s",
            len(file_bytes),
            exc,
        )
        raise HTTPException(
            status_code=400,
            detail="Could not read .docx file. Make sure it is a valid Word document (.docx, not .doc).",
        ) from exc

    # {normalized_label: [non-empty value lines]} from every two-column row.
    fields: dict[str, list[str]] = {}
    for table in document.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = _normalize_label(cells[0].text)
            if not label or label in fields:
                continue
            fields[label] = [
                p.text.strip() for p in cells[1].paragraphs if p.text.strip()
            ]

    def lines_for(*subs: str) -> list[str]:
        for norm_label, lines in fields.items():
            if all(sub in norm_label for sub in subs):
                return lines
        return []

    def build_items(lines, item_cls, id_attr, kind, qty_attr="quantity"):
        result = []
        for raw in _split_resource_cell(lines):
            parsed = _parse_resource_item_line(raw)
            if parsed is None:
                continue
            label, qty = parsed
            option = create_resource_option(db, kind, label, user_id)
            result.append(item_cls(**{id_attr: option.id, qty_attr: qty}))
        return result

    # "aircraft" must be matched exactly so it doesn't swallow the
    # "aircraft mission equipment" / "aircraft armament" rows.
    aircraft_lines = fields.get("aircraft", [])

    return ResourcesData(
        aircraftItems=build_items(
            aircraft_lines, AircraftItem, "aircraftTypeId", "aircraft_type",
            qty_attr="numberOfAircraft",
        ),
        aircraftMissionEquipmentItems=build_items(
            lines_for("aircraft", "mission", "equipment"),
            AircraftMissionEquipmentItem, "aircraftMissionEquipmentId",
            "aircraft_mission_equipment",
        ),
        aircraftArmamentItems=build_items(
            lines_for("aircraft", "armament"),
            AircraftArmamentItem, "aircraftArmamentId", "aircraft_armament",
        ),
        personalFlightEquipmentItems=build_items(
            lines_for("personal", "equipment"),
            PersonalFlightEquipmentItem, "personalFlightEquipmentId",
            "personal_flight_equipment",
        ),
        classroomRequirementItems=build_items(
            lines_for("classroom"),
            ClassroomRequirementItem, "classroomRequirementId",
            "classroom_requirements",
        ),
        trainingMaterialAidItems=build_items(
            lines_for("training material") or lines_for("aids"),
            TrainingMaterialAidItem, "trainingMaterialAidId",
            "training_material_aids",
        ),
    )


# ─── Lesson Planning .docx import (Instructional Scalar) ────────────────────


def _is_number_token(s: str) -> bool:
    """True for plain numbers like ``"2"`` or dotted refs like ``"1.10"``."""
    return bool(re.fullmatch(r"\d+(\.\d+)*", s))


def _first_text_from(cells: list[str], idx: int) -> str | None:
    """First non-empty, non-numeric cell at or after ``idx`` (the merged title)."""
    for j in range(idx, len(cells)):
        v = cells[j]
        if v and not _is_number_token(v):
            return v
    return None


def parse_lesson_planning_docx(
    db: Session, file_bytes: bytes, user_id: int
) -> LessonPlanningData:
    """Parse an Instructional Scalar .docx into Lesson Planning TO/EO/TP lists.

    The scalar is a fixed 10-column grid repeated across several tables. Per row:
      * Training Objective (TO): column 1 is an integer; the title spans col 2+.
      * Enabling Objective (EO): column 2 is a dotted number (e.g. "1.1"); the
        title follows the number columns.
      * Teaching Point (TP): the "Performance" column (index 5) statement.
    Only TO/EO/TP are imported; every other column is ignored. Each unique label
    becomes a master option via the idempotent :func:`create_resource_option`.
    """
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    try:
        import docx  # python-docx, imported lazily.

        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001 — any failure means a bad upload
        logger.warning(
            "Lesson Planning .docx import failed to parse (%d bytes): %s",
            len(file_bytes),
            exc,
        )
        raise HTTPException(
            status_code=400,
            detail="Could not read .docx file. Make sure it is a valid Word document (.docx, not .doc).",
        ) from exc

    to_labels: list[str] = []
    eo_labels: list[str] = []
    tp_labels: list[str] = []

    def add_unique(bucket: list[str], label: str | None) -> None:
        if label and label not in bucket:
            bucket.append(label)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 6:
                continue
            # Skip header / banner rows.
            if (
                "instructional scalar" in cells[0].lower()
                or cells[1] in ("TO", "Number")
                or cells[2] in ("EO", "Number")
            ):
                continue

            if cells[1] and re.fullmatch(r"\d+", cells[1]):
                add_unique(to_labels, _first_text_from(cells, 2))
            elif not cells[1] and cells[2] and "." in cells[2] and _is_number_token(cells[2]):
                add_unique(eo_labels, _first_text_from(cells, 3))
            elif not cells[1] and not cells[2]:
                # TP row: the Performance statement is the first text cell after
                # the Lsn/Ref (col 0) and the numeric TO/EO/TP columns. Its column
                # index shifts between tables — merged cells make the grid 9–11
                # wide — so scan for it instead of hard-coding index 5 (which lands
                # on the TP number or the Method "L"/"PI" column depending on table).
                add_unique(tp_labels, _first_text_from(cells, 1))

    def build(labels: list[str], kind: str, item_cls, id_attr: str):
        # Bulk-resolve to ids — these lists can be hundreds of items long.
        id_by_label = bulk_get_or_create_options(db, kind, labels, user_id)
        result = []
        for label in labels:
            opt_id = id_by_label.get(label.strip())
            if opt_id is not None:
                result.append(item_cls(**{id_attr: opt_id}))
        return result

    return LessonPlanningData(
        trainingObjectives=build(
            to_labels, "training_objective", TrainingObjectiveItem, "trainingObjectiveId"
        ),
        enablingObjectives=build(
            eo_labels, "enabling_objective", EnablingObjectiveItem, "enablingObjectiveId"
        ),
        teachingPoints=build(
            tp_labels, "teaching_point", TeachingPointItem, "teachingPointId"
        ),
    )


# ─── Lesson Creation .docx import (Detailed Syllabus + Instructional Scalar) ──


def _load_docx_or_400(file_bytes: bytes, what: str):
    """Open .docx bytes with python-docx, raising a 400 on any failure."""
    if not file_bytes:
        raise HTTPException(status_code=400, detail=f"{what} is empty.")
    try:
        import docx  # python-docx, imported lazily.

        return docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001 — any failure means a bad upload
        logger.warning("%s failed to parse (%d bytes): %s", what, len(file_bytes), exc)
        raise HTTPException(
            status_code=400,
            detail=f"Could not read {what}. Make sure it is a valid Word document (.docx, not .doc).",
        ) from exc


def _norm_lesson_code(s: str) -> str:
    """Normalize a lesson code for matching across the two documents.

    The Detailed Syllabus writes ``"AF1"`` / ``"SYS 2"`` while the Instructional
    Scalar writes ``"AF 1"`` / ``"SYS2"`` — collapsing whitespace and upper-casing
    makes them line up.
    """
    return re.sub(r"\s+", "", s or "").upper()


def parse_lesson_creation_docx(
    db: Session,
    syllabus_bytes: bytes,
    scalar_bytes: bytes,
    user_id: int,
) -> LessonCreationData:
    """Build Lesson Creation data from two documents.

    * Detailed Syllabus → the ordered lesson list with Period (Day/Night) and
      flight timing (number of periods).
    * Instructional Scalar → TO/EO/TP units attached to each lesson via the
      "Lsn / Ref" column (one ref may list several lesson codes, newline- or
      slash-separated, and one lesson may collect many units).
    """
    syllabus = _load_docx_or_400(syllabus_bytes, "Detailed Syllabus .docx")
    scalar = _load_docx_or_400(scalar_bytes, "Instructional Scalar .docx")

    # ── 1. Lessons from the Detailed Syllabus ────────────────────────────────
    # Each lesson: {number, title, period, timing, units:[(to, eo, tp)]}
    lessons: list[dict] = []
    by_code: dict[str, dict] = {}

    for table in syllabus.tables:
        # Locate the Day/Night header row to learn which columns are which.
        daynight: dict[int, str] = {}
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if "Day" in cells or "Night" in cells:
                daynight = {
                    i: v for i, v in enumerate(cells) if v in ("Day", "Night")
                }
                if daynight:
                    break

        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 7:
                continue
            number = cells[1]
            if (
                not number
                or number in ("Lesson Number", "Serial")
                or "DETAILED SYLLABUS" in cells[0]
            ):
                continue
            title = _first_text_from(cells, 2)
            if not title or title == "Lesson Title":
                continue

            day = night = 0
            for i, grp in daynight.items():
                if i < len(cells) and re.fullmatch(r"\d+", cells[i]):
                    if grp == "Day":
                        day += int(cells[i])
                    else:
                        night += int(cells[i])
            period = "Day" if (day > 0 and day >= night) else ("Night" if night > 0 else None)

            lesson = {
                "number": number,
                "title": title,
                "period": period,
                "timing": (day + night) or None,
                "units": [],
            }
            lessons.append(lesson)
            # First occurrence of a code wins for unit attachment (codes like
            # "AR1" repeat across phases; the scalar can't tell them apart).
            by_code.setdefault(_norm_lesson_code(number), lesson)

    # ── 2. TO/EO/TP units from the Instructional Scalar ──────────────────────
    cur_to: str | None = None
    cur_eo: str | None = None
    for table in scalar.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 6:
                continue
            if (
                "instructional scalar" in cells[0].lower()
                or cells[1] in ("TO", "Number")
                or cells[2] in ("EO", "Number")
            ):
                continue

            if cells[1] and re.fullmatch(r"\d+", cells[1]):
                cur_to = _first_text_from(cells, 2)
                cur_eo = None
            elif not cells[1] and cells[2] and "." in cells[2] and _is_number_token(cells[2]):
                cur_eo = _first_text_from(cells, 3)
            elif not cells[1] and not cells[2]:
                # See parse_lesson_planning_docx: the Performance statement's
                # column shifts between tables, so scan rather than read cells[5].
                tp = _first_text_from(cells, 1)
                ref = cells[0]
                if not tp or not ref:
                    continue
                for code in re.split(r"[/\n]", ref):
                    lesson = by_code.get(_norm_lesson_code(code))
                    if lesson is not None:
                        lesson["units"].append((cur_to, cur_eo, tp))

    # ── 3. Bulk-resolve every label to a master option id ────────────────────
    period_map = bulk_get_or_create_options(
        db, "type_of_period", [l["period"] for l in lessons if l["period"]], user_id
    )
    to_map = bulk_get_or_create_options(
        db, "training_objective",
        [u[0] for l in lessons for u in l["units"] if u[0]], user_id,
    )
    eo_map = bulk_get_or_create_options(
        db, "enabling_objective",
        [u[1] for l in lessons for u in l["units"] if u[1]], user_id,
    )
    tp_map = bulk_get_or_create_options(
        db, "teaching_point",
        [u[2] for l in lessons for u in l["units"] if u[2]], user_id,
    )

    def opt(mapping: dict[str, int], label: str | None) -> int | None:
        return mapping.get(label.strip()) if label else None

    # ── 4. Build the typed payload ───────────────────────────────────────────
    items: list[LessonItem] = []
    for lesson in lessons:
        units = [
            LessonUnitItem(
                trainingObjectiveId=opt(to_map, to),
                enablingObjectiveId=opt(eo_map, eo),
                teachingPointId=opt(tp_map, tp),
            )
            for (to, eo, tp) in lesson["units"]
        ]
        items.append(
            LessonItem(
                id=None,
                lessonNumber=lesson["number"],
                lessonTitle=lesson["title"],
                environmentId=None,
                periodTypeId=opt(period_map, lesson["period"]),
                flightTiming=lesson["timing"],
                units=units,
            )
        )

    return LessonCreationData(lessons=items)


def _sync_master_completion(master: CourseMaster) -> None:
    """Mirror the live completion into the CourseMaster's progress column so the
    Course Builder detail page reflects the change immediately.
    """
    master.course_info_completion = int(round(compute_completion_pct(master)))


# ─── Personnel Requirement tab ──────────────────────────────────────────────

def _personnel_to_data(row: CourseInfoPersonnelRequirement) -> PersonnelRequirementData:
    return PersonnelRequirementData(
        maximumTraineesNumber=row.maximum_trainees_number,
        minimumTraineesNumber=row.minimum_trainees_number,
        instructorsToTraineesRatioFlight=row.instructors_to_trainees_ratio_flight,
        instructorsToTraineesRatioClassroom=row.instructors_to_trainees_ratio_classroom,
        instructorsToTraineesRatioPractical=row.instructors_to_trainees_ratio_practical,
        instructorQualificationRequirements=[
            InstructorQualificationRequirementItem(
                instructorQualificationRequirementId=item.instructor_qualification_requirement_id
            )
            for item in row.instructor_qualification_requirements
        ],
    )


def _apply_personnel_data(
    db: Session, row, data: PersonnelRequirementData,
    *, item_orm_cls: Type = CourseInfoInstructorQualificationRequirement,
) -> None:
    row.maximum_trainees_number = data.maximumTraineesNumber
    row.minimum_trainees_number = data.minimumTraineesNumber
    row.instructors_to_trainees_ratio_flight = data.instructorsToTraineesRatioFlight
    row.instructors_to_trainees_ratio_classroom = data.instructorsToTraineesRatioClassroom
    row.instructors_to_trainees_ratio_practical = data.instructorsToTraineesRatioPractical
    _rebuild_label_only_item_list(
        db,
        collection=row.instructor_qualification_requirements,
        items=data.instructorQualificationRequirements or [],
        master_model=MasterInstructorQualificationRequirement,
        item_orm_cls=item_orm_cls,
        fk_attr="instructor_qualification_requirement_id",
        id_attr="instructorQualificationRequirementId",
    )


def get_personnel_requirement_tab(
    master: CourseMaster,
) -> tuple[str, PersonnelRequirementData | None]:
    """Return ``(status, data)`` for the Personnel Requirement tab."""
    row = master.personnel_requirement
    if row is None:
        return "incomplete", None
    return row.status, _personnel_to_data(row)


def upsert_personnel_requirement_tab(
    db: Session,
    master: CourseMaster,
    new_status: str,
    data: PersonnelRequirementData,
    user_id: int,
) -> CourseInfoPersonnelRequirement:
    if new_status not in ("incomplete", "draft", "complete"):
        raise HTTPException(
            status_code=400,
            detail="status must be 'incomplete', 'draft' or 'complete'",
        )

    row = master.personnel_requirement
    _guard_unmark_when_approved(master, row, new_status)
    if row is None:
        row = CourseInfoPersonnelRequirement(
            status=new_status,
            created_by_id=user_id,
            updated_by_id=user_id,
        )
        _apply_personnel_data(db, row, data)
        # Use the relationship so master.personnel_requirement reflects
        # the new row before we compute the completion percentage below.
        master.personnel_requirement = row
    else:
        row.status = new_status
        _apply_personnel_data(db, row, data)
        row.updated_by_id = user_id

    _sync_master_completion(master)

    db.commit()
    db.refresh(row)
    return row


# ─── Lesson Planning tab ────────────────────────────────────────────────────

def _lesson_planning_to_data(row: CourseInfoLessonPlanning) -> LessonPlanningData:
    return LessonPlanningData(
        numberOfPeriodsPerDay=row.number_of_periods_per_day,
        numberOfTrainingDaysPerWeek=row.number_of_training_days_per_week,
        numberOfPeriodsPerHalfDay=row.number_of_periods_per_half_day,
        periodDurationMinutes=row.period_duration_minutes,
        trainingObjectives=[TrainingObjectiveItem(trainingObjectiveId=t.training_objective_id) for t in row.training_objectives],
        enablingObjectives=[EnablingObjectiveItem(enablingObjectiveId=e.enabling_objective_id) for e in row.enabling_objectives],
        teachingPoints=[TeachingPointItem(teachingPointId=p.teaching_point_id) for p in row.teaching_points],
        typeOfPeriods=[TypeOfPeriodItem(typeOfPeriodId=tp.type_of_period_id) for tp in row.type_of_periods],
        classificationOfPeriods=[ClassificationOfPeriodItem(classificationOfPeriodId=cp.classification_of_period_id) for cp in row.classification_of_periods],
        environments=[EnvironmentItem(environmentId=e.environment_id) for e in row.environments],
    )


# Default child ORM classes for the Lesson Planning collections. The Course
# Selection instance category passes its own mirror classes (same shape) to
# reuse ``_apply_lesson_planning_data``.
_LESSON_PLANNING_ITEM_CLASSES: dict[str, Type] = {
    "training_objectives": CourseInfoLessonPlanningTrainingObjective,
    "enabling_objectives": CourseInfoLessonPlanningEnablingObjective,
    "teaching_points": CourseInfoLessonPlanningTeachingPoint,
    "type_of_periods": CourseInfoLessonPlanningTypeOfPeriod,
    "classification_of_periods": CourseInfoLessonPlanningClassificationOfPeriod,
    "environments": CourseInfoLessonPlanningEnvironment,
}


def _apply_lesson_planning_data(
    db: Session, row, data: LessonPlanningData,
    *, classes: dict[str, Type] | None = None,
) -> None:
    cls = classes or _LESSON_PLANNING_ITEM_CLASSES
    row.number_of_periods_per_day = data.numberOfPeriodsPerDay
    row.number_of_training_days_per_week = data.numberOfTrainingDaysPerWeek
    row.number_of_periods_per_half_day = data.numberOfPeriodsPerHalfDay
    row.period_duration_minutes = data.periodDurationMinutes

    _rebuild_label_only_item_list(
        db,
        collection=row.training_objectives,
        items=data.trainingObjectives or [],
        master_model=MasterTrainingObjective,
        item_orm_cls=cls["training_objectives"],
        fk_attr="training_objective_id",
        id_attr="trainingObjectiveId",
    )
    _rebuild_label_only_item_list(
        db,
        collection=row.enabling_objectives,
        items=data.enablingObjectives or [],
        master_model=MasterEnablingObjective,
        item_orm_cls=cls["enabling_objectives"],
        fk_attr="enabling_objective_id",
        id_attr="enablingObjectiveId",
    )
    _rebuild_label_only_item_list(
        db,
        collection=row.teaching_points,
        items=data.teachingPoints or [],
        master_model=MasterTeachingPoint,
        item_orm_cls=cls["teaching_points"],
        fk_attr="teaching_point_id",
        id_attr="teachingPointId",
    )
    _rebuild_label_only_item_list(
        db,
        collection=row.type_of_periods,
        items=data.typeOfPeriods or [],
        master_model=MasterPeriodType,
        item_orm_cls=cls["type_of_periods"],
        fk_attr="type_of_period_id",
        id_attr="typeOfPeriodId",
    )
    _rebuild_label_only_item_list(
        db,
        collection=row.classification_of_periods,
        items=data.classificationOfPeriods or [],
        master_model=MasterPeriodClassification,
        item_orm_cls=cls["classification_of_periods"],
        fk_attr="classification_of_period_id",
        id_attr="classificationOfPeriodId",
    )
    _rebuild_label_only_item_list(
        db,
        collection=row.environments,
        items=data.environments or [],
        master_model=MasterEnvironment,
        item_orm_cls=cls["environments"],
        fk_attr="environment_id",
        id_attr="environmentId",
    )


def get_lesson_planning_tab(
    master: CourseMaster,
) -> tuple[str, LessonPlanningData | None]:
    """Return ``(status, data)`` for the Lesson Planning tab."""
    row = master.lesson_planning
    if row is None:
        return "incomplete", None
    return row.status, _lesson_planning_to_data(row)


def upsert_lesson_planning_tab(
    db: Session,
    master: CourseMaster,
    new_status: str,
    data: LessonPlanningData,
    user_id: int,
) -> CourseInfoLessonPlanning:
    if new_status not in ("incomplete", "draft", "complete"):
        raise HTTPException(
            status_code=400,
            detail="status must be 'incomplete', 'draft' or 'complete'",
        )

    row = master.lesson_planning
    _guard_unmark_when_approved(master, row, new_status)
    if row is None:
        row = CourseInfoLessonPlanning(
            status=new_status,
            created_by_id=user_id,
            updated_by_id=user_id,
        )
        db.add(row)  # ← must be in session for cascade to work
        _apply_lesson_planning_data(db, row, data)
        master.lesson_planning = row
    else:
        row.status = new_status
        _apply_lesson_planning_data(db, row, data)
        row.updated_by_id = user_id

    _sync_master_completion(master)

    db.commit()
    db.refresh(row)
    return row


# ─── Lesson Creation tab ────────────────────────────────────────────────────

def _lesson_creation_to_data(row: CourseInfoLessonCreation) -> LessonCreationData:
    return LessonCreationData(
        lessons=[
            LessonItem(
                id=lesson.id,
                lessonNumber=lesson.lesson_number,
                lessonTitle=lesson.lesson_title,
                environmentId=lesson.environment_id,
                periodTypeId=lesson.period_type_id,
                flightTiming=lesson.flight_timing,
                periodPerUnit=lesson.period_per_unit,
                instructorStudentRatio=lesson.instructor_student_ratio,
                location=lesson.location,
                healthAndSafety=lesson.health_and_safety,
                units=[
                    LessonUnitItem(
                        trainingObjectiveId=u.training_objective_id,
                        enablingObjectiveId=u.enabling_objective_id,
                        teachingPointId=u.teaching_point_id,
                    )
                    for u in lesson.units
                ],
                resources=[
                    LessonResourceItem(
                        category=r.category,
                        resourceId=r.resource_id,
                    )
                    for r in lesson.resources
                ],
                conducts=[
                    LessonConductItem(
                        part=c.part,
                        point=c.point,
                        material=c.material,
                        notes=c.notes,
                    )
                    for c in lesson.conducts
                ],
            )
            for lesson in row.lessons
        ]
    )


def _apply_lesson_creation_data(
    db: Session, row, data: LessonCreationData,
    *,
    lesson_cls: Type = CourseInfoLessonCreationLesson,
    unit_cls: Type = CourseInfoLessonCreationLessonUnit,
    resource_cls: Type = CourseInfoLessonCreationLessonResource,
    conduct_cls: Type = CourseInfoLessonCreationLessonConduct,
) -> None:
    cleaned: list[LessonItem] = []
    for item in data.lessons:
        units = item.units or []
        non_blank_units = [
            u for u in units
            if u.trainingObjectiveId is not None
            or u.enablingObjectiveId is not None
            or u.teachingPointId is not None
        ]
        # A resource only counts once both the category and the picked
        # resource are present — a half-filled row is dropped.
        non_blank_resources = [
            r for r in (item.resources or [])
            if r.category is not None and r.resourceId is not None
        ]
        # A conduct record needs a part and at least one of point/material/notes.
        non_blank_conducts = [
            c for c in (item.conducts or [])
            if c.part is not None
            and (c.point is not None or c.material is not None or c.notes is not None)
        ]
        # Drop fully-blank rows (basics + units + resources + conducts all empty)
        if (
            item.lessonNumber is None
            and item.lessonTitle is None
            and item.environmentId is None
            and item.periodTypeId is None
            and item.flightTiming is None
            and item.instructorStudentRatio is None
            and item.location is None
            and item.healthAndSafety is None
            and not non_blank_units
            and not non_blank_resources
            and not non_blank_conducts
        ):
            continue
        _ensure_master_row_exists(db, MasterEnvironment, item.environmentId)
        _ensure_master_row_exists(db, MasterPeriodType, item.periodTypeId)
        for u in non_blank_units:
            _ensure_master_row_exists(db, MasterTrainingObjective, u.trainingObjectiveId)
            _ensure_master_row_exists(db, MasterEnablingObjective, u.enablingObjectiveId)
            _ensure_master_row_exists(db, MasterTeachingPoint, u.teachingPointId)
        for r in non_blank_resources:
            # _resolve_master_model 400s on an unknown category, then we
            # confirm the chosen resource row actually exists in that table.
            _ensure_master_row_exists(db, _resolve_master_model(r.category), r.resourceId)
        item = item.model_copy(
            update={
                "units": non_blank_units,
                "resources": non_blank_resources,
                "conducts": non_blank_conducts,
            }
        )
        cleaned.append(item)

    # Diff-upsert keyed by lesson id so existing lesson rows keep their PK
    # across edits. Material files/folders FK to these PKs with ON DELETE
    # CASCADE, so a wipe-and-recreate here would silently delete all uploaded
    # material on every save — hence the in-place update below.
    existing = {lesson.id: lesson for lesson in row.lessons}
    kept_ids: set[int] = set()
    for idx, item in enumerate(cleaned):
        lesson_row = existing.get(item.id) if item.id is not None else None
        if lesson_row is None:
            lesson_row = lesson_cls()
            row.lessons.append(lesson_row)
        else:
            kept_ids.add(lesson_row.id)
        lesson_row.order_index = idx
        lesson_row.lesson_number = item.lessonNumber
        lesson_row.lesson_title = item.lessonTitle
        lesson_row.environment_id = item.environmentId
        lesson_row.period_type_id = item.periodTypeId
        lesson_row.flight_timing = item.flightTiming
        lesson_row.period_per_unit = item.periodPerUnit
        lesson_row.instructor_student_ratio = item.instructorStudentRatio
        lesson_row.location = item.location
        lesson_row.health_and_safety = item.healthAndSafety
        # Units aren't linked to material, so clear-and-recreate is safe.
        lesson_row.units.clear()
        for u_idx, u in enumerate(item.units or []):
            lesson_row.units.append(
                unit_cls(
                    order_index=u_idx,
                    training_objective_id=u.trainingObjectiveId,
                    enabling_objective_id=u.enablingObjectiveId,
                    teaching_point_id=u.teachingPointId,
                )
            )
        # Resources are likewise standalone — clear-and-recreate.
        lesson_row.resources.clear()
        for r_idx, r in enumerate(item.resources or []):
            lesson_row.resources.append(
                resource_cls(
                    order_index=r_idx,
                    category=r.category,
                    resource_id=r.resourceId,
                )
            )
        # Conduct records — clear-and-recreate, order preserved within parts.
        lesson_row.conducts.clear()
        for c_idx, c in enumerate(item.conducts or []):
            lesson_row.conducts.append(
                conduct_cls(
                    order_index=c_idx,
                    part=c.part,
                    point=c.point,
                    material=c.material,
                    notes=c.notes,
                )
            )

    # Remove lesson rows the client dropped (delete-orphan cascades units +
    # cascades any material linked to those lessons).
    for lesson in list(row.lessons):
        if lesson.id is not None and lesson.id not in kept_ids:
            row.lessons.remove(lesson)


def get_lesson_creation_tab(
    master: CourseMaster,
) -> tuple[str, LessonCreationData | None]:
    """Return ``(status, data)`` for the Lesson Creation tab."""
    row = master.lesson_creation
    if row is None:
        return "incomplete", None
    return row.status, _lesson_creation_to_data(row)


def get_lesson_associations(
    db: Session, master: CourseMaster, lesson_id: int
) -> LessonAssociationsResponse:
    """Count how many Material, Form Builder, Evaluation and Schedule records
    reference the given Lesson Creation lesson.

    The four downstream categories attach their content to a lesson by FK with
    ``ON DELETE CASCADE``, so removing a lesson silently deletes everything
    counted here. The Lesson Creation tab calls this before a delete so it can
    warn the user when the lesson is still in use. 404s if the lesson does not
    belong to this course master.
    """
    # Local imports keep the cross-module model dependencies out of import time
    # (these modules don't import course_info.services, so this stays acyclic).
    from app.modules.evaluation.models import (
        EvaluationLessonForm,
        EvaluationLessonQuiz,
    )
    from app.modules.form_builder.models import (
        FormBuilderFormLink,
        FormBuilderSurveyLink,
    )
    from app.modules.material.models import MaterialFile, MaterialFolder
    from app.modules.schedule.models import (
        CourseScheduleDay,
        CourseSchedulePlacement,
    )

    valid_ids = {lesson.id for lesson in (master.lesson_creation.lessons if master.lesson_creation else [])}
    if lesson_id not in valid_ids:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Lesson not found for this course",
        )

    def _count(model, *extra_where) -> int:
        return int(
            db.execute(
                select(func.count(model.id)).where(
                    model.lesson_id == lesson_id, *extra_where
                )
            ).scalar_one()
            or 0
        )

    material = _count(MaterialFile) + _count(MaterialFolder)
    form_builder = _count(FormBuilderSurveyLink) + _count(FormBuilderFormLink)
    evaluation = _count(EvaluationLessonQuiz) + _count(EvaluationLessonForm)
    schedule = int(
        db.execute(
            select(func.count(CourseSchedulePlacement.id))
            .join(
                CourseScheduleDay,
                CourseSchedulePlacement.course_schedule_day_id == CourseScheduleDay.id,
            )
            .where(
                CourseScheduleDay.course_master_id == master.id,
                CourseSchedulePlacement.lesson_id == lesson_id,
            )
        ).scalar_one()
        or 0
    )

    return LessonAssociationsResponse(
        material=material,
        formBuilder=form_builder,
        evaluation=evaluation,
        schedule=schedule,
        total=material + form_builder + evaluation + schedule,
    )


# Per-lesson resource picklist is sourced from the Resources tab. Each entry:
#   (CourseInfoResources relationship, child-item master relationship,
#    combo-box kind stored on the lesson resource, human category label).
_LESSON_RESOURCE_CATEGORIES: tuple[tuple[str, str, str, str], ...] = (
    ("aircraft_items", "aircraft_type", "aircraft_type", "Aircraft"),
    ("aircraft_mission_equipment_items", "aircraft_mission_equipment", "aircraft_mission_equipment", "Aircraft Mission Equipment"),
    ("aircraft_armament_items", "aircraft_armament", "aircraft_armament", "Aircraft Armament"),
    ("ground_equipment_simulator_type_items", "simulator_type", "ground_equipment_simulator_type", "Simulator Type"),
    ("ground_equipment_mission_planning_system_items", "mission_planning_system", "ground_equipment_mission_planning_system", "Mission Planning System"),
    ("ground_equipment_maintenance_items", "ground_maintenance_equipment", "ground_equipment_maintenance", "Ground Maintenance Equipment"),
    ("ground_equipment_armament_items", "ground_armament", "ground_equipment_armament", "Ground Armament"),
    ("personal_flight_equipment_items", "personal_flight_equipment", "personal_flight_equipment", "Personal Flight Equipment"),
    ("aviation_life_support_equipment_items", "aviation_life_support_equipment", "aviation_life_support_equipment", "Aviation Life Support Equipment"),
    ("classroom_requirement_items", "classroom_requirement", "classroom_requirements", "Classroom Requirements"),
    ("training_material_aid_items", "training_material_aid", "training_material_aids", "Training Material Aids"),
)


def list_lesson_resource_options(
    master: CourseMaster,
) -> list[LessonResourceCategoryOption]:
    """Resources available to a lesson, grouped by category (master owner)."""
    return list_lesson_resource_options_for(master.resources)


def list_lesson_resource_options_for(resources) -> list[LessonResourceCategoryOption]:
    """Resources available to a lesson, grouped by category.

    Sourced from the given Resources tab row — only categories that have at
    least one resource added there are returned, each carrying that course's
    distinct resources (deduped by id, sorted by label) for the second combo.
    Owner-agnostic: works for both the master row and the Course Selection
    instance row (same child relationship attribute names).
    """
    result: list[LessonResourceCategoryOption] = []
    if resources is None:
        return result
    for rel_attr, master_attr, kind, label in _LESSON_RESOURCE_CATEGORIES:
        seen: dict[int, str] = {}
        for item in getattr(resources, rel_attr, []) or []:
            option_master = getattr(item, master_attr, None)
            if option_master is not None:
                seen.setdefault(option_master.id, option_master.label)
        if not seen:
            continue
        options = [
            ResourceOptionResponse(id=rid, label=lbl)
            for rid, lbl in sorted(seen.items(), key=lambda kv: kv[1].lower())
        ]
        result.append(
            LessonResourceCategoryOption(
                category=kind, categoryLabel=label, resources=options
            )
        )
    return result


# ─── TP Link Preview tab (derived, read-only) ───────────────────────────────

def get_tp_link_preview(db: Session, master: CourseMaster) -> TpLinkPreviewData:
    """TP Link Preview for a master owner (see :func:`get_tp_link_preview_for`)."""
    return get_tp_link_preview_for(db, master.lesson_creation)


def get_tp_link_preview_for(db: Session, lesson_creation) -> TpLinkPreviewData:
    """Build the TO -> EO -> TP tree of teaching points used by lessons, plus
    the list of master teaching points not yet associated with any lesson.

    A teaching point is "associated" when a lesson unit references it; the unit
    also supplies the TO/EO it sits under. Units missing a TO or EO are bucketed
    under an "Unassigned …" node so their teaching point still appears. Owner-
    agnostic: ``lesson_creation`` may be the master or instance row.
    """
    complete = bool(lesson_creation and lesson_creation.status == "complete")

    associated_tp_ids: set[int] = set()
    # Nested insertion-ordered dicts: TO id -> EO id -> {tp_id: tp_label}.
    # ``None`` ids are valid keys (Unassigned buckets).
    tos: dict = {}
    if lesson_creation is not None:
        for lesson in lesson_creation.lessons:
            for u in lesson.units:
                if u.teaching_point_id is None:
                    continue
                associated_tp_ids.add(u.teaching_point_id)
                to_label = (
                    u.training_objective.label
                    if u.training_objective
                    else "Unassigned Training Objective"
                )
                eo_label = (
                    u.enabling_objective.label
                    if u.enabling_objective
                    else "Unassigned Enabling Objective"
                )
                tp_label = (
                    u.teaching_point.label
                    if u.teaching_point
                    else f"Teaching Point #{u.teaching_point_id}"
                )
                to_entry = tos.setdefault(
                    u.training_objective_id,
                    {"id": u.training_objective_id, "label": to_label, "eos": {}},
                )
                eo_entry = to_entry["eos"].setdefault(
                    u.enabling_objective_id,
                    {"id": u.enabling_objective_id, "label": eo_label, "tps": {}},
                )
                eo_entry["tps"].setdefault(u.teaching_point_id, tp_label)

    def _by_label(items):
        return sorted(items, key=lambda e: (e["label"] or "").lower())

    associated: list[TpLinkTrainingObjective] = []
    for to_entry in _by_label(tos.values()):
        eos: list[TpLinkEnablingObjective] = []
        for eo_entry in _by_label(to_entry["eos"].values()):
            tps = [
                TpLinkTeachingPoint(id=tid, label=lbl)
                for tid, lbl in sorted(
                    eo_entry["tps"].items(), key=lambda kv: (kv[1] or "").lower()
                )
            ]
            eos.append(
                TpLinkEnablingObjective(
                    enablingObjectiveId=eo_entry["id"],
                    label=eo_entry["label"],
                    teachingPoints=tps,
                )
            )
        associated.append(
            TpLinkTrainingObjective(
                trainingObjectiveId=to_entry["id"],
                label=to_entry["label"],
                enablingObjectives=eos,
            )
        )

    all_tps = db.scalars(
        select(MasterTeachingPoint).order_by(MasterTeachingPoint.label)
    ).all()
    unassociated = [
        TpLinkTeachingPoint(id=tp.id, label=tp.label)
        for tp in all_tps
        if tp.id not in associated_tp_ids
    ]

    return TpLinkPreviewData(
        associated=associated,
        unassociated=unassociated,
        lessonCreationComplete=complete,
    )


def upsert_lesson_creation_tab(
    db: Session,
    master: CourseMaster,
    new_status: str,
    data: LessonCreationData,
    user_id: int,
) -> CourseInfoLessonCreation:
    if new_status not in ("incomplete", "draft", "complete"):
        raise HTTPException(
            status_code=400,
            detail="status must be 'incomplete', 'draft' or 'complete'",
        )

    row = master.lesson_creation
    _guard_unmark_when_approved(master, row, new_status)
    if row is None:
        row = CourseInfoLessonCreation(
            status=new_status,
            created_by_id=user_id,
            updated_by_id=user_id,
        )
        db.add(row)
        _apply_lesson_creation_data(db, row, data)
        master.lesson_creation = row
    else:
        row.status = new_status
        _apply_lesson_creation_data(db, row, data)
        row.updated_by_id = user_id

    _sync_master_completion(master)

    db.commit()
    db.refresh(row)
    return row
